"""Generador de Word (.docx) con todos los campos extraídos."""
from __future__ import annotations

import io

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor

from app.models.document import Document, DocType
from app.models.extraction import Extraction


def generate_word(doc: Document, extraction: Extraction) -> io.BytesIO:
    buf = io.BytesIO()
    d = DocxDocument()

    if doc.doc_type == DocType.INVOICE and extraction:
        _write_invoice(d, extraction)
    elif doc.doc_type == DocType.CONTRACT and extraction:
        _write_contract(d, extraction)
    elif doc.doc_type == DocType.TABLE and extraction:
        _write_table(d, extraction)
    else:
        _write_generic(d, doc, extraction)

    d.save(buf)
    buf.seek(0)
    return buf


def _f(fields: dict, key: str):
    return fields.get(key, {}).get("value")


def _add_field(p, label: str, value) -> None:
    if value is None or value == "":
        return
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(str(value)).font.size = Pt(10)


def _add_section_heading(d: DocxDocument, title: str) -> None:
    h = d.add_heading(title, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(29, 78, 216)


def _write_invoice(d: DocxDocument, ext: Extraction) -> None:
    fields = ext.data.get("fields", {})
    items = ext.data.get("items", [])
    ibp = ext.data.get("ibp_entries", [])
    tasas = ext.data.get("tasa_vial_entries", [])
    riesgos = ext.data.get("risk_descriptions", [])

    # Título
    inv_num = _f(fields, "invoice_number") or ""
    inv_letter = _f(fields, "invoice_letter") or ""
    d.add_heading(f"Factura {inv_letter} N° {inv_num}", level=0)

    # ── Emisor ──
    _add_section_heading(d, "Emisor")
    for key, label in [
        ("emitter_name", "Razón Social"), ("emitter_cuit", "CUIT"),
        ("emitter_address", "Domicilio"), ("emitter_iibb", "IIBB"),
        ("emitter_iva_condition", "Condición IVA"), ("emitter_start_date", "Inicio Actividades"),
    ]:
        p = d.add_paragraph()
        _add_field(p, label, _f(fields, key))

    # ── Receptor ──
    _add_section_heading(d, "Receptor")
    for key, label in [
        ("receptor_name", "Razón Social"), ("receptor_cuit", "CUIT"),
        ("receptor_address", "Domicilio"), ("receptor_iibb", "IIBB"),
        ("receptor_account", "N° Cuenta"), ("receptor_deudor_account", "Cuenta Deudora"),
    ]:
        p = d.add_paragraph()
        _add_field(p, label, _f(fields, key))

    # ── Datos Factura ──
    _add_section_heading(d, "Datos de la Factura")
    for key, label in [
        ("emission_date", "Fecha"), ("cae", "CAE"),
        ("invoice_cae", "CAE (pie)"), ("cae_expiry", "Vencimiento CAE"),
        ("incoterms", "Incoterms"), ("sap_number", "N° SAP"),
        ("oc_number", "OC"), ("payment_terms", "Condiciones de Pago"),
        ("shipping_method", "Forma de Envío"),
    ]:
        p = d.add_paragraph()
        _add_field(p, label, _f(fields, key))

    # ── Ítems ──
    if items:
        _add_section_heading(d, f"Ítems ({len(items)})")
        headers = ["Código", "Descripción", "Cantidad", "UM", "Riesgo ONU", "V. Unitario", "Importe"]
        table = d.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for item in items:
            row = table.add_row().cells
            row[0].text = str(item.get("code", ""))
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("quantity", ""))
            row[3].text = str(item.get("unit", ""))
            row[4].text = str(item.get("risk_un", ""))
            row[5].text = str(item.get("unit_price", ""))
            row[6].text = str(item.get("import", ""))
        d.add_paragraph("")

    # ── Totales y Desglose de Resumen ──
    _add_section_heading(d, "Totales y Desglose de Impuestos")
    for key, label in [
        ("subtotal", "Subtotal"), ("net", "Neto Gravado"), ("importe_neto", "Importe Neto"),
        ("iva_amount", "IVA"), ("iva_inscripto", "IVA Inscripto"), ("iva_no_inscripto", "IVA No Inscripto"),
        ("iibb", "IIBB"), ("ingresos_brutos", "Ing. Brutos"),
        ("tasas_municipales", "Tasas Municipales"), ("sellos", "Sellos"),
        ("percepcion_iva", "Percepción IVA"), ("iva_percepcion", "IVA Percepción"),
        ("itc", "ITC (Combustibles)"), ("co2", "CO2 (Dióxido de Carbono)"),
        ("financiacion", "Financiación"), ("icl_amount", "ICL"), ("idc_amount", "IDC"),
        ("iva_percentage", "% IVA"), ("tasa_vial", "Total Tasa Vial"),
        ("total", "TOTAL"),
    ]:
        p = d.add_paragraph()
        _add_field(p, label, _f(fields, key))

    summary_breakdown = ext.data.get("summary_breakdown", [])
    if summary_breakdown:
        _add_section_heading(d, "Resumen de Impuestos y Totales")
        t = d.add_table(rows=1, cols=2, style="Light Grid Accent 1")
        t.rows[0].cells[0].text = "Concepto / Impuesto"
        t.rows[0].cells[1].text = "Monto"
        for entry in summary_breakdown:
            row = t.add_row().cells
            row[0].text = str(entry.get("label", ""))
            row[1].text = str(entry.get("amount", ""))
        d.add_paragraph("")

    # ── IBP ──
    if ibp:
        _add_section_heading(d, "Ingresos Brutos Provinciales (IBP)")
        t = d.add_table(rows=1, cols=3, style="Light Grid Accent 1")
        for i, h in enumerate(["Jurisdicción", "Alícuota %", "Importe"]):
            t.rows[0].cells[i].text = h
        for entry in ibp:
            row = t.add_row().cells
            row[0].text = str(entry.get("jurisdiction", ""))
            row[1].text = str(entry.get("percentage", ""))
            row[2].text = str(entry.get("amount", ""))
        d.add_paragraph("")

    # ── Tasa Vial ──
    if tasas:
        _add_section_heading(d, "Tasa Vial por Localidad")
        t = d.add_table(rows=1, cols=2, style="Light Grid Accent 1")
        t.rows[0].cells[0].text = "Localidad"
        t.rows[0].cells[1].text = "Importe"
        for entry in tasas:
            row = t.add_row().cells
            row[0].text = str(entry.get("locality", ""))
            row[1].text = str(entry.get("amount", ""))
        d.add_paragraph("")

    # ── Riesgo/ONU ──
    if riesgos:
        _add_section_heading(d, "Riesgo / ONU")
        for entry in riesgos:
            p = d.add_paragraph()
            run = p.add_run(f"N° {entry.get('number', '')}: ")
            run.bold = True
            p.add_run(str(entry.get("description", "")))


def _write_contract(d: DocxDocument, ext: Extraction) -> None:
    meta = ext.data.get("meta", {})
    fields = ext.data.get("fields", {})
    title = fields.get("title", {}).get("value", "Documento")
    d.add_heading(title, level=0)
    full_text = meta.get("full_text", "")
    for line in full_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            d.add_paragraph("")
        elif stripped.isupper() and len(stripped) < 80:
            d.add_heading(stripped, level=2)
        else:
            d.add_paragraph(stripped)


def _write_table(d: DocxDocument, ext: Extraction) -> None:
    tables = ext.data.get("tables", [])
    for tbl_data in tables:
        rows = tbl_data.get("rows", [])
        if not rows:
            continue
        table = d.add_table(rows=len(rows), cols=len(rows[0]), style="Light Grid Accent 1")
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = str(val or "")
        d.add_paragraph("")


def _write_generic(d: DocxDocument, doc: Document | None, ext: Extraction | None) -> None:
    d.add_heading("Documento", level=0)
    if ext:
        meta = ext.data.get("meta", {})
        full_text = meta.get("full_text", "")
        if full_text:
            for line in full_text.split("\n"):
                d.add_paragraph(line)
