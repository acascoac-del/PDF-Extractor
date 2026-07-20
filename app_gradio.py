"""PDF Extractor - App Gradio para Hugging Face Spaces.

Aplicación de un solo archivo que reutiliza los módulos de extracción existentes
en app/services/ y proporciona una interfaz web simple para extraer datos de PDFs.

Tabs:
  1. Extraer PDF — carga, extracción, vista de resultados, descarga.
  2. Configuración IA — proveedor, API key, modelo.
  3. Acerca de — descripción, limitaciones, link de suscripción.
"""
from __future__ import annotations

import json
import logging
import os
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

import gradio as gr

# ---------------------------------------------------------------------------
# Asegurar que el directorio raíz del proyecto esté en sys.path para imports
# ---------------------------------------------------------------------------
_APP_ROOT = Path(__file__).resolve().parent
if str(_APP_ROOT) not in sys.path:
    sys.path.insert(0, str(_APP_ROOT))

# ---------------------------------------------------------------------------
# Imports de los módulos existentes (sin dependencia a DB)
# ---------------------------------------------------------------------------
try:
    from app.services.pdf_text import extract_text, PdfContent
    from app.services.extraction.base import ExtractionResult
    from app.services.extraction.invoice import extract_invoice
    from app.services.extraction.table import extract_table
    from app.services.extraction.contract import extract_contract
    from app.services.extraction.generic import extract_generic
    from app.services.classification import classify_by_rules
    from app.services.confidence import compute_overall_confidence
    from app.services.llm import get_client, extract_invoice_with_llm
except ImportError as exc:
    logging.getLogger(__name__).warning(
        "No se pudieron importar módulos de app/: %s. "
        "La extracción no estará disponible.", exc
    )
    # Definir stubs para que el módulo se pueda cargar y mostrar error amigable.
    extract_text = None  # type: ignore[assignment]
    PdfContent = None  # type: ignore[assignment,misc]
    ExtractionResult = None  # type: ignore[assignment,misc]
    extract_invoice = None  # type: ignore[assignment]
    extract_table = None  # type: ignore[assignment]
    extract_contract = None  # type: ignore[assignment]
    extract_generic = None  # type: ignore[assignment]
    classify_by_rules = None  # type: ignore[assignment]
    compute_overall_confidence = None  # type: ignore[assignment]
    get_client = None  # type: ignore[assignment]
    extract_invoice_with_llm = None  # type: ignore[assignment]

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------
USAGE_FILE = _APP_ROOT / "usage.json"
SETTINGS_FILE = _APP_ROOT / "user_settings.json"
FREE_LIMIT = 3
PRO_PRICE = "$5 USD/mes"
MERCADO_PAGO_LINK = "https://mpago.la/2QkZ4XJ"  # Reemplazar con link real

DOC_TYPE_LABELS = {
    "Auto-detectar": "auto",
    "Factura": "invoice",
    "Tabla": "table",
    "Contrato": "contract",
    "Generico": "generic",
}

PROVIDER_DEFAULTS = {
    "Ninguno (solo reglas)": {"base_url": "", "model": ""},
    "OpenAI": {"base_url": "https://api.openai.com/v1", "model": "gpt-4o-mini"},
    "OpenRouter": {"base_url": "https://openrouter.ai/api/v1", "model": "openai/gpt-4o-mini"},
    "Groq": {"base_url": "https://api.groq.com/openai/v1", "model": "llama-3.1-8b-instant"},
    "Ollama": {"base_url": "http://localhost:11434/v1", "model": "llama3.1"},
    "Personalizado": {"base_url": "", "model": ""},
}

# ---------------------------------------------------------------------------
# Archivos de persistencia (JSON simple)
# ---------------------------------------------------------------------------


def _load_json(path: Path, default: dict) -> dict:
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            pass
    return default.copy()


def _save_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_usage() -> dict:
    current_month = datetime.now().strftime("%Y-%m")
    usage = _load_json(USAGE_FILE, {"count": 0, "month": current_month, "pro": False})
    # Resetear contador si cambió el mes
    if usage.get("month") != current_month:
        usage = {"count": 0, "month": current_month, "pro": usage.get("pro", False)}
        _save_json(USAGE_FILE, usage)
    return usage


def save_usage(usage: dict) -> None:
    _save_json(USAGE_FILE, usage)


def load_settings() -> dict:
    return _load_json(SETTINGS_FILE, {
        "provider": "Ninguno (solo reglas)",
        "api_key": "",
        "base_url": "",
        "model": "",
    })


def save_settings(settings: dict) -> None:
    _save_json(SETTINGS_FILE, settings)


# ---------------------------------------------------------------------------
# Adaptadores ligeros para exportadores (evitar import de ORM)
# ---------------------------------------------------------------------------


class _MockDoc:
    """Simula app.models.document.Document para los exportadores."""

    def __init__(self, doc_type: str, filename: str = ""):
        self.doc_type = doc_type
        self.original_filename = filename


class _MockExtraction:
    """Simula app.models.extraction.Extraction para los exportadores."""

    def __init__(self, data: dict, doc_type: str = "generic", llm_model: str | None = None):
        self.data = data
        self.doc_type = doc_type
        self.llm_model = llm_model


def _generate_excel_standalone(doc_type: str, data: dict) -> bytes:
    """Genera Excel sin depender de los modelos ORM."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    fields = data.get("fields", {})
    items = data.get("items", [])
    ibp = data.get("ibp_entries", [])
    tasas = data.get("tasa_vial_entries", [])
    riesgos = data.get("risk_descriptions", [])
    tables = data.get("tables", [])

    hfont = Font(bold=True, size=11, color="FFFFFF")
    sfont = Font(bold=True, size=12, color="1D4ED8")
    hfill = PatternFill(start_color="1D4ED8", end_color="1D4ED8", fill_type="solid")
    lfill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
    thin = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )

    def _set_header(ws, row, headers):
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font = hfont
            c.fill = hfill
            c.alignment = Alignment(horizontal="center")
            c.border = thin

    def _section_row(ws, row, title, cols=2):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
        c = ws.cell(row=row, column=1, value=title)
        c.font = sfont
        c.fill = lfill
        return row + 1

    def _field_row(ws, row, label, value):
        if value is None or value == "":
            return row
        ws.cell(row=row, column=1, value=label).font = Font(bold=True, size=10)
        ws.cell(row=row, column=1).border = thin
        c = ws.cell(row=row, column=2, value=value)
        c.border = thin
        c.alignment = Alignment(wrap_text=True)
        return row + 1

    def _f(key):
        return fields.get(key, {}).get("value")

    if doc_type == "invoice":
        ws = wb.active
        ws.title = "Detalle"
        r = 1
        # Emisor
        r = _section_row(ws, r, "EMISOR")
        for k, lbl in [("emitter_name", "Razon Social"), ("emitter_cuit", "CUIT"),
                        ("emitter_address", "Domicilio"), ("emitter_iibb", "IIBB"),
                        ("emitter_iva_condition", "Condicion IVA"), ("emitter_start_date", "Inicio Actividades")]:
            r = _field_row(ws, r, lbl, _f(k))
        # Receptor
        r = _section_row(ws, r, "RECEPTOR")
        for k, lbl in [("receptor_name", "Razon Social"), ("receptor_cuit", "CUIT"),
                        ("receptor_address", "Domicilio"), ("receptor_iibb", "IIBB"),
                        ("receptor_account", "N Cuenta"), ("receptor_deudor_account", "Cuenta Deudora")]:
            r = _field_row(ws, r, lbl, _f(k))
        # Datos factura
        r = _section_row(ws, r, "DATOS DE LA FACTURA")
        for k, lbl in [("invoice_number", "N Factura"), ("invoice_letter", "Tipo (A/B/C)"),
                        ("emission_date", "Fecha Emision"), ("cae", "CAE"),
                        ("cae_expiry", "Vencimiento CAE"), ("incoterms", "Incoterms"),
                        ("sap_number", "N SAP"), ("oc_number", "OC"),
                        ("payment_terms", "Condiciones Pago")]:
            r = _field_row(ws, r, lbl, _f(k))
        # Totales
        r = _section_row(ws, r, "TOTALES")
        for k, lbl in [("importe_neto", "Importe Neto"), ("financiacion", "Financiacion"),
                        ("iva_inscripto", "IVA Inscripto"), ("iva_percentage", "% IVA"),
                        ("ingresos_brutos", "Ing. Brutos"), ("tasa_vial", "Tasa Vial"),
                        ("net", "Neto Gravado"), ("iva_amount", "IVA"), ("total", "TOTAL")]:
            r = _field_row(ws, r, lbl, _f(k))
        # IBP
        if ibp:
            r = _section_row(ws, r, "IBP", 3)
            _set_header(ws, r, ["Jurisdiccion", "Alicuota %", "Importe"])
            r += 1
            for entry in ibp:
                ws.cell(row=r, column=1, value=entry.get("jurisdiction", "")).border = thin
                ws.cell(row=r, column=2, value=entry.get("percentage")).border = thin
                ws.cell(row=r, column=3, value=entry.get("amount")).border = thin
                r += 1
        # Tasa Vial
        if tasas:
            r = _section_row(ws, r, "TASA VIAL", 2)
            _set_header(ws, r, ["Localidad", "Importe"])
            r += 1
            for entry in tasas:
                ws.cell(row=r, column=1, value=entry.get("locality", "")).border = thin
                ws.cell(row=r, column=2, value=entry.get("amount")).border = thin
                r += 1
        # Riesgo
        if riesgos:
            r = _section_row(ws, r, "RIESGO / ONU", 2)
            _set_header(ws, r, ["N Riesgo", "Descripcion"])
            r += 1
            for entry in riesgos:
                ws.cell(row=r, column=1, value=entry.get("number", "")).border = thin
                ws.cell(row=r, column=2, value=entry.get("description", "")).border = thin
                r += 1
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 45
        # Items
        if items:
            ws2 = wb.create_sheet("Items")
            hdrs = ["Codigo", "Descripcion", "Cantidad", "UM", "Riesgo ONU",
                    "V. Unitario", "Importe"]
            _set_header(ws2, 1, hdrs)
            for i, item in enumerate(items, 2):
                ws2.cell(row=i, column=1, value=item.get("code", "")).border = thin
                ws2.cell(row=i, column=2, value=item.get("description", "")).border = thin
                ws2.cell(row=i, column=3, value=item.get("quantity")).border = thin
                ws2.cell(row=i, column=4, value=item.get("unit", "")).border = thin
                ws2.cell(row=i, column=5, value=item.get("risk_un")).border = thin
                ws2.cell(row=i, column=6, value=item.get("unit_price")).border = thin
                ws2.cell(row=i, column=7, value=item.get("import")).border = thin
            for c, w in enumerate([12, 25, 14, 8, 14, 16, 18], 1):
                ws2.column_dimensions[get_column_letter(c)].width = w
        # Resumen
        ws3 = wb.create_sheet("Resumen")
        shdrs = ["N Factura", "Tipo", "Emisor", "CUIT", "Receptor", "CUIT Receptor",
                 "Fecha", "CAE", "TOTAL"]
        _set_header(ws3, 1, shdrs)
        row_data = [_f("invoice_number"), _f("invoice_letter"), _f("emitter_name"),
                     _f("emitter_cuit"), _f("receptor_name"), _f("receptor_cuit"),
                     _f("emission_date"), _f("cae"), _f("total")]
        for col, val in enumerate(row_data, 1):
            ws3.cell(row=2, column=col, value=val).border = thin

    elif doc_type == "table" and tables:
        for i, tbl in enumerate(tables):
            ws = wb.active if i == 0 else wb.create_sheet(f"Tabla {i + 1}")
            if i == 0:
                ws.title = "Tabla 1"
            rows = tbl.get("rows", [])
            if rows:
                _set_header(ws, 1, rows[0])
                for r_idx, row in enumerate(rows[1:], 2):
                    for c_idx, val in enumerate(row, 1):
                        ws.cell(row=r_idx, column=c_idx, value=val).border = thin
                for c in range(1, len(rows[0]) + 1):
                    ws.column_dimensions[get_column_letter(c)].width = 18

    elif doc_type == "contract":
        ws = wb.active
        ws.title = "Texto"
        title = fields.get("title", {}).get("value", "Documento")
        ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=14)
        meta = data.get("meta", {})
        full_text = meta.get("full_text", "")
        for i, line in enumerate(full_text.split("\n"), 3):
            ws.cell(row=i, column=1, value=line)
        ws.column_dimensions["A"].width = 100

    else:
        ws = wb.active
        ws.title = "Datos"
        ws.cell(row=1, column=1, value="Campo").font = Font(bold=True)
        ws.cell(row=1, column=2, value="Valor").font = Font(bold=True)
        row = 2
        for key, f in fields.items():
            ws.cell(row=row, column=1, value=key)
            ws.cell(row=row, column=2, value=f.get("value"))
            row += 1
        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 60

    import io
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generate_word_standalone(doc_type: str, data: dict) -> bytes:
    """Genera Word sin depender de los modelos ORM."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    import io

    d = DocxDocument()
    fields = data.get("fields", {})
    items = data.get("items", [])
    ibp = data.get("ibp_entries", [])
    tasas = data.get("tasa_vial_entries", [])
    riesgos = data.get("risk_descriptions", [])

    def _f(key):
        return fields.get(key, {}).get("value")

    def _add_field(p, label, value):
        if value is None or value == "":
            return
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(str(value)).font.size = Pt(10)

    def _add_heading(title):
        h = d.add_heading(title, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(29, 78, 216)

    if doc_type == "invoice":
        inv_num = _f("invoice_number") or ""
        inv_letter = _f("invoice_letter") or ""
        d.add_heading(f"Factura {inv_letter} N {inv_num}", level=0)

        _add_heading("Emisor")
        for k, lbl in [("emitter_name", "Razon Social"), ("emitter_cuit", "CUIT"),
                        ("emitter_address", "Domicilio")]:
            p = d.add_paragraph()
            _add_field(p, lbl, _f(k))

        _add_heading("Receptor")
        for k, lbl in [("receptor_name", "Razon Social"), ("receptor_cuit", "CUIT"),
                        ("receptor_address", "Domicilio")]:
            p = d.add_paragraph()
            _add_field(p, lbl, _f(k))

        _add_heading("Datos Factura")
        for k, lbl in [("emission_date", "Fecha"), ("cae", "CAE"),
                        ("cae_expiry", "Vto. CAE"), ("incoterms", "Incoterms"),
                        ("sap_number", "SAP"), ("oc_number", "OC"),
                        ("payment_terms", "Condiciones Pago")]:
            p = d.add_paragraph()
            _add_field(p, lbl, _f(k))

        if items:
            _add_heading(f"Items ({len(items)})")
            hdrs = ["Codigo", "Descripcion", "Cantidad", "UM", "Riesgo ONU", "V. Unitario", "Importe"]
            table = d.add_table(rows=1, cols=len(hdrs), style="Light Grid Accent 1")
            for i, h in enumerate(hdrs):
                table.rows[0].cells[i].text = h
            for item in items:
                row = table.add_row().cells
                row[0].text = str(item.get("code", ""))
                row[1].text = str(item.get("description", ""))
                row[2].text = str(item.get("quantity", ""))
                row[3].text = str(item.get("unit", ""))
                row[4].text = str(item.get("risk_un", ""))
                row[5].text = str(item.get("unit_price", ""))
                row[6].text = str(item.get("import", ""))

        _add_heading("Totales")
        for k, lbl in [("importe_neto", "Importe Neto"), ("iva_inscripto", "IVA Inscripto"),
                        ("iva_percentage", "% IVA"), ("total", "TOTAL")]:
            p = d.add_paragraph()
            _add_field(p, lbl, _f(k))

        if ibp:
            _add_heading("Ingresos Brutos Provinciales (IBP)")
            t = d.add_table(rows=1, cols=3, style="Light Grid Accent 1")
            for i, h in enumerate(["Jurisdiccion", "Alicuota %", "Importe"]):
                t.rows[0].cells[i].text = h
            for entry in ibp:
                row = t.add_row().cells
                row[0].text = str(entry.get("jurisdiction", ""))
                row[1].text = str(entry.get("percentage", ""))
                row[2].text = str(entry.get("amount", ""))

        if tasas:
            _add_heading("Tasa Vial por Localidad")
            t = d.add_table(rows=1, cols=2, style="Light Grid Accent 1")
            t.rows[0].cells[0].text = "Localidad"
            t.rows[0].cells[1].text = "Importe"
            for entry in tasas:
                row = t.add_row().cells
                row[0].text = str(entry.get("locality", ""))
                row[1].text = str(entry.get("amount", ""))

        if riesgos:
            _add_heading("Riesgo / ONU")
            for entry in riesgos:
                p = d.add_paragraph()
                run = p.add_run(f"N {entry.get('number', '')}: ")
                run.bold = True
                p.add_run(str(entry.get("description", "")))

    elif doc_type == "contract":
        meta = data.get("meta", {})
        title = fields.get("title", {}).get("value", "Documento")
        d.add_heading(title, level=0)
        full_text = meta.get("full_text", "")
        for line in full_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                d.add_paragraph("")
            elif stripped.isupper() and len(stripped) < 80:
                d.add_heading(stripped, level=2)
            else:
                d.add_paragraph(stripped)

    elif doc_type == "table":
        tables = data.get("tables", [])
        for tbl_data in tables:
            rows = tbl_data.get("rows", [])
            if rows:
                table = d.add_table(rows=len(rows), cols=len(rows[0]), style="Light Grid Accent 1")
                for r_idx, row in enumerate(rows):
                    for c_idx, val in enumerate(row):
                        table.rows[r_idx].cells[c_idx].text = str(val or "")
                d.add_paragraph("")

    else:
        d.add_heading("Documento", level=0)
        meta = data.get("meta", {})
        full_text = meta.get("full_text", "")
        if full_text:
            for line in full_text.split("\n"):
                d.add_paragraph(line)
        for key, f in fields.items():
            p = d.add_paragraph()
            _add_field(p, key, f.get("value"))

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generate_csv_standalone(data: dict) -> str:
    """Genera CSV sin depender de los modelos ORM."""
    import csv
    import io

    buf = io.StringIO()
    fields = data.get("fields", {})
    items = data.get("items", [])

    writer = csv.writer(buf, delimiter=";", quoting=csv.QUOTE_ALL)
    writer.writerow(["Campo", "Valor", "Confianza", "Fuente"])
    for key, f in fields.items():
        writer.writerow([key, f.get("value", ""), f.get("confidence", ""), f.get("source", "")])

    if items:
        writer.writerow([])
        writer.writerow(["# Item", "Descripcion", "Cantidad", "Precio Unit.", "Subtotal"])
        for i, item in enumerate(items, 1):
            writer.writerow([
                i,
                item.get("description", ""),
                item.get("quantity", ""),
                item.get("unit_price", ""),
                item.get("subtotal") or item.get("import", ""),
            ])

    buf.seek(0)
    return buf.getvalue()


def _generate_json_standalone(data: dict) -> bytes:
    """Genera JSON sin depender de los modelos ORM."""
    import io

    def _clean(d):
        if isinstance(d, dict):
            return {k: _clean(v) for k, v in d.items()}
        if isinstance(d, list):
            return [_clean(i) if isinstance(i, dict) else i for i in d]
        return d

    buf = io.BytesIO()
    buf.write(json.dumps(_clean(data), ensure_ascii=False, indent=2, default=str).encode("utf-8"))
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Funcion principal de extraccion
# ---------------------------------------------------------------------------


def _check_modules():
    """Verifica que los modulos de extraccion esten disponibles."""
    if extract_text is None:
        raise RuntimeError(
            "Los modulos de extraccion no se pudieron importar. "
            "Verifique que la carpeta app/ este completa."
        )


def extract_pdf(file_path: str, doc_type_label: str) -> tuple:
    """Procesa un PDF y devuelve todos los resultados.

    Returns:
        Tupla con: (status_msg, summary_md, fields_df, items_df,
                     ibp_df, tasa_vial_df, risk_df,
                     excel_path, word_path, csv_path, json_path)
    """
    try:
        _check_modules()
    except RuntimeError as exc:
        empty_df = gr.update(value=[], visible=False)
        return (
            f"**Error:** {exc}", "", empty_df, empty_df, empty_df, empty_df, empty_df,
            None, None, None, None,
        )

    # Verificar limite de uso
    usage = load_usage()
    if not usage.get("pro", False) and usage.get("count", 0) >= FREE_LIMIT:
        msg = (
            f"**Limite alcanzado** ({FREE_LIMIT} PDFs/mes). "
            f"Suscribite por {PRO_PRICE} para uso ilimitado.\n\n"
            f"[Suscribirse en Mercado Pago]({MERCADO_PAGO_LINK})"
        )
        empty_df = gr.update(value=[], visible=False)
        return (msg, "", empty_df, empty_df, empty_df, empty_df, empty_df,
                None, None, None, None)

    if not file_path:
        empty_df = gr.update(value=[], visible=False)
        return (
            "**Error:** No se selecciono ningun archivo.", "", empty_df, empty_df,
            empty_df, empty_df, empty_df, None, None, None, None,
        )

    # Validar tamano (30 MB)
    try:
        file_size = os.path.getsize(file_path)
        if file_size > 30 * 1024 * 1024:
            empty_df = gr.update(value=[], visible=False)
            return (
                "**Error:** El archivo supera los 30 MB permitidos.", "", empty_df,
                empty_df, empty_df, empty_df, empty_df, None, None, None, None,
            )
    except OSError:
        pass

    try:
        # 1. Extraer texto
        content: PdfContent = extract_text(file_path)

        # 2. Clasificar
        doc_type_key = DOC_TYPE_LABELS.get(doc_type_label, "auto")
        if doc_type_key == "auto":
            doc_type, cls_confidence = classify_by_rules(content.full_text)
            doc_type_str = doc_type.value if hasattr(doc_type, "value") else str(doc_type)
        else:
            doc_type_str = doc_type_key
            cls_confidence = 1.0

        # 3. Configurar LLM si esta disponible
        llm_client = None
        settings = load_settings()
        if settings.get("api_key") and settings.get("base_url"):
            try:
                from openai import OpenAI
                llm_client = OpenAI(
                    base_url=settings["base_url"],
                    api_key=settings["api_key"],
                    timeout=60,
                )
            except Exception:
                pass  # LLM es opcional

        # 4. Extraer datos segun tipo
        result: ExtractionResult
        if doc_type_str == "invoice":
            result = extract_invoice(content, llm_client=llm_client)
        elif doc_type_str == "table":
            result = extract_table(content)
        elif doc_type_str == "contract":
            result = extract_contract(content)
        else:
            result = extract_generic(content)

        # 5. Calcular confianza
        overall_conf = compute_overall_confidence(result.data)

        # 6. Preparar datos para UI
        data = result.data
        fields = data.get("fields", {})
        items = data.get("items", [])
        ibp_entries = data.get("ibp_entries", [])
        tasa_vial_entries = data.get("tasa_vial_entries", [])
        risk_descs = data.get("risk_descriptions", [])
        meta = data.get("meta", {})

        # Summary
        type_display = {
            "invoice": "Factura",
            "table": "Tabla",
            "contract": "Contrato",
            "generic": "Generico",
        }.get(doc_type_str, doc_type_str)

        scanned_str = "Si" if meta.get("is_scanned") else "No"
        summary_md = (
            f"**Tipo de documento:** {type_display} "
            f"(confianza clasificacion: {cls_confidence:.0%})\n\n"
            f"**Confianza general de extraccion:** {overall_conf:.0%}\n\n"
            f"**Paginas:** {content.page_count} | "
            f"**Escaneado (OCR):** {scanned_str}"
        )

        # Fields table
        fields_rows = []
        for key, f in fields.items():
            val = f.get("value")
            if val is None or val == "":
                continue
            source = f.get("source", "")
            conf = f.get("confidence", 0.0)
            if isinstance(conf, (int, float)):
                if conf >= 0.8:
                    badge = "Alta"
                elif conf >= 0.5:
                    badge = "Media"
                else:
                    badge = "Baja"
            else:
                badge = str(conf)
            fields_rows.append([key, str(val), source, badge])

        # Items table
        items_rows = []
        for item in items:
            items_rows.append([
                str(item.get("code", "")),
                str(item.get("description", "")),
                str(item.get("quantity", "")),
                str(item.get("unit", "")),
                str(item.get("risk_un", "")),
                str(item.get("unit_price", "")),
                str(item.get("import") or item.get("subtotal", "")),
            ])

        # IBP table
        ibp_rows = []
        for entry in ibp_entries:
            ibp_rows.append([
                str(entry.get("jurisdiction", "")),
                str(entry.get("percentage", "")),
                str(entry.get("amount", "")),
            ])

        # Tasa Vial table
        tasa_rows = []
        for entry in tasa_vial_entries:
            tasa_rows.append([
                str(entry.get("locality", "")),
                str(entry.get("amount", "")),
            ])

        # Risk table
        risk_rows = []
        for entry in risk_descs:
            risk_rows.append([
                str(entry.get("number", "")),
                str(entry.get("description", "")),
            ])

        # 7. Generar archivos de exportacion
        tmp_dir = tempfile.mkdtemp()
        base_name = Path(file_path).stem

        excel_path = os.path.join(tmp_dir, f"{base_name}.xlsx")
        word_path = os.path.join(tmp_dir, f"{base_name}.docx")
        csv_path = os.path.join(tmp_dir, f"{base_name}.csv")
        json_path = os.path.join(tmp_dir, f"{base_name}.json")

        try:
            with open(excel_path, "wb") as f:
                f.write(_generate_excel_standalone(doc_type_str, data))
        except Exception as exc:
            logger.warning("Error generando Excel: %s", exc)
            excel_path = None

        try:
            with open(word_path, "wb") as f:
                f.write(_generate_word_standalone(doc_type_str, data))
        except Exception as exc:
            logger.warning("Error generando Word: %s", exc)
            word_path = None

        try:
            with open(csv_path, "w", encoding="utf-8") as f:
                f.write(_generate_csv_standalone(data))
        except Exception as exc:
            logger.warning("Error generando CSV: %s", exc)
            csv_path = None

        try:
            with open(json_path, "wb") as f:
                f.write(_generate_json_standalone(data))
        except Exception as exc:
            logger.warning("Error generando JSON: %s", exc)
            json_path = None

        # 8. Incrementar uso
        usage["count"] = usage.get("count", 0) + 1
        save_usage(usage)

        # 9. Preparar visibilidad de tablas
        has_items = len(items_rows) > 0
        has_ibp = len(ibp_rows) > 0
        has_tasa = len(tasa_rows) > 0
        has_risk = len(risk_rows) > 0

        status_msg = f"**Extraccion completada exitosamente.** Archivos generados para descarga."

        return (
            status_msg,
            summary_md,
            gr.update(value=fields_rows, visible=True),
            gr.update(value=items_rows, visible=has_items),
            gr.update(value=ibp_rows, visible=has_ibp),
            gr.update(value=tasa_rows, visible=has_tasa),
            gr.update(value=risk_rows, visible=has_risk),
            excel_path,
            word_path,
            csv_path,
            json_path,
        )

    except Exception as exc:
        logger.exception("Error en extraccion")
        empty_df = gr.update(value=[], visible=False)
        error_msg = f"**Error al procesar el PDF:** {str(exc)}"
        return (
            error_msg, "", empty_df, empty_df, empty_df, empty_df, empty_df,
            None, None, None, None,
        )


# ---------------------------------------------------------------------------
# Configuracion IA
# ---------------------------------------------------------------------------


def on_provider_change(provider: str) -> tuple:
    """Actualiza base_url y model cuando cambia el proveedor."""
    defaults = PROVIDER_DEFAULTS.get(provider, {"base_url": "", "model": ""})
    return defaults["base_url"], defaults["model"]


def test_connection(api_key: str, base_url: str, model: str) -> str:
    """Prueba la conexion con el proveedor LLM."""
    if not api_key or not base_url:
        return "**Error:** Ingrese API key y URL base."
    if not model:
        return "**Error:** Ingrese un nombre de modelo."

    try:
        from openai import OpenAI
        client = OpenAI(base_url=base_url, api_key=api_key, timeout=15)
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": "Responde solo: OK"}],
            max_tokens=5,
        )
        reply = resp.choices[0].message.content.strip()
        return f"**Conexion exitosa.** Respuesta del modelo: {reply}"
    except Exception as exc:
        return f"**Error de conexion:** {str(exc)}"


def save_ai_config(provider: str, api_key: str, base_url: str, model: str) -> str:
    """Guarda la configuracion de IA en user_settings.json."""
    settings = {
        "provider": provider,
        "api_key": api_key,
        "base_url": base_url,
        "model": model,
    }
    save_settings(settings)
    return "**Configuracion guardada correctamente.**"


# ---------------------------------------------------------------------------
# Construccion de la interfaz Gradio
# ---------------------------------------------------------------------------


def build_app() -> gr.Blocks:
    """Construye y devuelve la aplicacion Gradio."""

    theme = gr.themes.Soft(
        primary_hue="blue",
        secondary_hue="green",
    )

    with gr.Blocks(
        theme=theme,
        title="PDF Extractor - Extraccion Inteligente",
        css="""
            .main-header { text-align: center; margin-bottom: 1em; }
            .main-header h1 { margin-bottom: 0.2em; }
            .download-section { border: 1px solid #e0e0e0; border-radius: 8px; padding: 1em; margin-top: 1em; }
        """,
    ) as app:

        # ---- Header ----
        gr.Markdown(
            """
            <div class="main-header">
                <h1>PDF Extractor</h1>
                <p>Extraccion inteligente de facturas, tablas y contratos</p>
            </div>
            """,
        )

        with gr.Tabs():
            # ================================================================
            # TAB 1: Extraer PDF
            # ================================================================
            with gr.Tab("Extraer PDF"):
                with gr.Row():
                    with gr.Column(scale=1):
                        file_input = gr.File(
                            label="Subir PDF (max 30 MB)",
                            file_types=[".pdf"],
                            type="filepath",
                        )
                        doc_type_dropdown = gr.Dropdown(
                            choices=list(DOC_TYPE_LABELS.keys()),
                            value="Auto-detectar",
                            label="Tipo de documento",
                        )
                        extract_btn = gr.Button(
                            "Extraer",
                            variant="primary",
                            size="lg",
                        )

                    with gr.Column(scale=2):
                        status_msg = gr.Markdown(label="Estado")
                        summary_md = gr.Markdown(label="Resumen")

                # Tablas de resultados
                fields_df = gr.Dataframe(
                    headers=["Campo", "Valor", "Fuente", "Confianza"],
                    label="Campos extraidos",
                    visible=True,
                    interactive=False,
                )

                with gr.Row():
                    items_df = gr.Dataframe(
                        headers=["Codigo", "Descripcion", "Cantidad", "UM",
                                 "Riesgo ONU", "V. Unitario", "Importe"],
                        label="Items (factura)",
                        visible=False,
                        interactive=False,
                    )

                with gr.Row():
                    ibp_df = gr.Dataframe(
                        headers=["Jurisdiccion", "Alicuota %", "Importe"],
                        label="Ingresos Brutos Provinciales (IBP)",
                        visible=False,
                        interactive=False,
                    )
                    tasa_vial_df = gr.Dataframe(
                        headers=["Localidad", "Importe"],
                        label="Tasa Vial por Localidad",
                        visible=False,
                        interactive=False,
                    )

                risk_df = gr.Dataframe(
                    headers=["N Riesgo", "Descripcion"],
                    label="Riesgo / ONU",
                    visible=False,
                    interactive=False,
                )

                # Descargas
                gr.Markdown("### Descargar resultados")
                with gr.Row():
                    with gr.Column():
                        excel_file = gr.File(label="Excel (.xlsx)", interactive=True)
                    with gr.Column():
                        word_file = gr.File(label="Word (.docx)", interactive=True)
                    with gr.Column():
                        csv_file = gr.File(label="CSV (.csv)", interactive=True)
                    with gr.Column():
                        json_file = gr.File(label="JSON (.json)", interactive=True)

                # Evento de extraccion
                extract_btn.click(
                    fn=extract_pdf,
                    inputs=[file_input, doc_type_dropdown],
                    outputs=[
                        status_msg, summary_md,
                        fields_df, items_df,
                        ibp_df, tasa_vial_df, risk_df,
                        excel_file, word_file, csv_file, json_file,
                    ],
                )

            # ================================================================
            # TAB 2: Configuracion IA
            # ================================================================
            with gr.Tab("Configuracion IA"):
                gr.Markdown(
                    """
                    ### Configuracion del motor de IA

                    El extractor funciona con reglas por defecto (sin IA).
                    Podes configurar un proveedor LLM para mejorar la extraccion
                    de campos complejos.
                    """
                )

                # Cargar configuracion guardada
                saved = load_settings()

                provider_dd = gr.Dropdown(
                    choices=list(PROVIDER_DEFAULTS.keys()),
                    value=saved.get("provider", "Ninguno (solo reglas)"),
                    label="Proveedor",
                )
                api_key_input = gr.Textbox(
                    label="API Key",
                    type="password",
                    value=saved.get("api_key", ""),
                    placeholder="sk-...",
                )
                base_url_input = gr.Textbox(
                    label="URL Base",
                    value=saved.get("base_url", ""),
                    placeholder="https://api.openai.com/v1",
                )
                model_input = gr.Textbox(
                    label="Modelo",
                    value=saved.get("model", ""),
                    placeholder="gpt-4o-mini",
                )

                with gr.Row():
                    test_btn = gr.Button("Probar conexion", variant="secondary")
                    save_btn = gr.Button("Guardar configuracion", variant="primary")

                config_status = gr.Markdown("")

                # Auto-completar URL y modelo al cambiar proveedor
                provider_dd.change(
                    fn=on_provider_change,
                    inputs=[provider_dd],
                    outputs=[base_url_input, model_input],
                )

                test_btn.click(
                    fn=test_connection,
                    inputs=[api_key_input, base_url_input, model_input],
                    outputs=[config_status],
                )

                save_btn.click(
                    fn=save_ai_config,
                    inputs=[provider_dd, api_key_input, base_url_input, model_input],
                    outputs=[config_status],
                )

            # ================================================================
            # TAB 3: Acerca de
            # ================================================================
            with gr.Tab("Acerca de"):
                gr.Markdown(
                    f"""
                    ## PDF Extractor - Extraccion Inteligente

                    Herramienta para extraer datos estructurados de documentos PDF
                    (facturas, tablas, contratos) y exportarlos a Excel, Word, CSV y JSON.

                    ### Como funciona

                    1. **Subis un PDF** con el boton de carga.
                    2. **El sistema detecta** automaticamente el tipo de documento
                       (factura, tabla, contrato) o podes seleccionarlo manualmente.
                    3. **Extrae los campos** usando un motor hibrido de reglas + IA.
                    4. **Descargas** los resultados en el formato que prefieras.

                    ### Tipos de documento soportados

                    | Tipo | Descripcion |
                    |------|-------------|
                    | **Factura** | Facturas AFIP (A/B/C): CUIT, CAE, items, totales |
                    | **Tabla** | Tablas genericas: reconstruccion de filas y columnas |
                    | **Contrato** | Contratos y documentos de texto plano |
                    | **Generico** | Cualquier otro documento PDF |

                    ### Motor hibrido (reglas + IA)

                    - **Reglas**: Extrae campos comunes (CUIT, CAE, fechas, montos)
                      con expresiones regulares. No requiere conexion externa.
                    - **IA** (opcional): Configura un proveedor LLM (OpenAI, Groq,
                      OpenRouter, Ollama) para mejorar la extraccion de campos
                      complejos como descripciones de items o formas de pago.

                    ### Limitaciones

                    - Archivos de hasta **30 MB**.
                    - **3 extracciones gratuitas por mes** (se renueva cada mes).
                    - El OCR requiere Tesseract instalado (para PDFs escaneados).
                    - La precision varia segun la calidad del PDF.
                    - No se garantiza 100% de exactitud en los datos extraidos.

                    ### Uso gratuito y Pro

                    - **Gratis**: 3 PDFs por mes.
                    - **Pro** ({PRO_PRICE}): Uso ilimitado, prioridad en el procesamiento.

                    [Suscribirse en Mercado Pago]({MERCADO_PAGO_LINK})

                    ---

                    *Desarrollado con Gradio + Python. Motor de extraccion: regex + OpenAI-compatible API.*
                    """
                )

    return app


# ---------------------------------------------------------------------------
# Punto de entrada
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    app = build_app()
    app.launch()
